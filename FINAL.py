import sys
from PySide6.QtWidgets import *
from PySide6.QtCore import *
from PySide6.QtGui import *
from MAINWINDOW_ui import Ui_GRAIN_ANALYSIS_SOFTWARE
from secondwindow_ui import Ui_secondwindow
import cv2, csv, numpy as np
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas, NavigationToolbar2QT
from matplotlib.figure import Figure
from sklearn.linear_model import LinearRegression
import pytesseract, re, os, pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from io import BytesIO
import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


class ImageProcessor:
    """Class to handle all image processing operations"""
    def __init__(self):
        # Get the root Application directory from UI folder
        current_dir = os.path.dirname(os.path.abspath(__file__))  # Gets UI folder
        application_dir = os.path.dirname(os.path.dirname(current_dir))  # Goes up two levels
        
        if os.name == 'nt':  # Windows
            tesseract_exe = 'tesseract.exe'
        else:  # Linux/Mac
            tesseract_exe = 'tesseract.exe'
            
        # Build path to Tesseract in Application/vendors
        tesseract_path = os.path.join(
            application_dir,
            'vendors',
            'tesseract',
            tesseract_exe
        )
        
        # Verify Tesseract exists
        if not os.path.exists(tesseract_path):
            raise FileNotFoundError(
                f"Tesseract not found at: {tesseract_path}\n"
                "Please ensure Tesseract is installed in the vendors folder"
            )
            
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Set up tessdata path
        tessdata_path = os.path.join(
            application_dir,
            'vendors',
            'tesseract',
            'tessdata'
        )
        
        if not os.path.exists(tessdata_path):
            raise FileNotFoundError(
                f"Tessdata folder not found at: {tessdata_path}\n"
                "Please ensure language data is installed"
            )
            
        os.environ['TESSDATA_PREFIX'] = tessdata_path
            
    def extract_metadata(self, image):
        """Extract time and temperature from image"""
        height, width = image.shape[:2]
        top_left = image[0:height//4, 0:width//4]
        enhanced = self.enhance_text_region(top_left)
        
        # Extract text from image
        text = pytesseract.image_to_string(
            enhanced,
            config='--psm 6 --oem 3'
        )
        
        # Extract time
        time_match = re.search(r'(\d+\.?\d*)\s*sec', text, re.IGNORECASE)
        time = float(time_match.group(1)) if time_match else 0.0
        
        # Extract temperature
        temp_match = re.search(r'(\d+\.?\d*)\s*[C°]', text, re.IGNORECASE)
        temp = float(temp_match.group(1)) if temp_match else 0.0
        
        return time, temp
 
    def adjust_brightness(self , image, brightness=30):
       
       hsv_image = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
       h, s, v = cv2.split(hsv_image)
       lim = 255 - brightness
       v[v > lim] = 255
       v[v <= lim] += brightness
       brightened_image = cv2.merge((h, s, v))
       brightened_image = cv2.cvtColor(brightened_image, cv2.COLOR_HSV2BGR)
       return brightened_image
 
    def enhance_text_region(self, image):
        """Enhance image for better text extraction"""
        
        # Enhance contrast
        gray =cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        enhanced = cv2.equalizeHist(gray)
        return enhanced

    def process_image(self, image_path,threshold=30):
      # Read image
      grain_sizes_count = []
      grain_count = 0
      image = cv2.imread(image_path)
      if image is None:
        raise ValueError(f"Could not read image: {image_path}")
      processed = self.preprocess_for_grain_analysis(image)
      mask, center, radius = self.create_circular_mask(processed)
      contours = self.detect_grains(processed, mask)
      gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
      _, dark_mask = cv2.threshold(gray_image, threshold, 255, cv2.THRESH_BINARY_INV)
      for contour in contours:
        mask_contour = np.zeros_like(gray_image, dtype=np.uint8)
        cv2.drawContours(mask_contour, [contour], -1, 255, thickness=cv2.FILLED)
        combined_mask = cv2.bitwise_or(mask_contour, dark_mask)
        area = cv2.countNonZero(combined_mask)
        if area <50 :  # Filter out noise
                continue
        # Calculate centroid
        M = cv2.moments(contour)
        if M["m00"] == 0:
            continue
        
        cx = int(M["m10"] / M["m00"])
        cy = int(M["m01"] / M["m00"])
        
        # Check if grain is within mask
        distance_to_center = np.sqrt((cx - center[0]) ** 2 + (cy - center[1]) ** 2)
        if distance_to_center + np.sqrt(area / np.pi) <= radius:
            grain_count += 1  # Fully inside the circle
        elif distance_to_center <= radius:
            grain_count += 0.5  # Partially within the circle  
        grain_count= int(grain_count)
            
            
            
            
      # Extract metadata
      time, temp = self.extract_metadata(image)
      brightened_image = self.adjust_brightness(image)

      # Process image for grain analysis
      processed = self.preprocess_for_grain_analysis(brightened_image)
      mask, center, radius = self.create_circular_mask(processed)
      contours = self.detect_grains(processed, mask)

      gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
      _, dark_mask = cv2.threshold(gray_image, threshold, 255, cv2.THRESH_BINARY_INV)
    
      # Loop through contours and analyze grains
      for contour in contours:
        # Calculate area
        mask_contour = np.zeros_like(gray_image, dtype=np.uint8)
        cv2.drawContours(mask_contour, [contour], -1, 255, thickness=cv2.FILLED)
        combined_mask = cv2.bitwise_or(mask_contour, dark_mask)
        area = cv2.countNonZero(combined_mask)
        grain_sizes_count.append(area)
        if area < 50:  # Filter out noise
            continue
  

        
        
        
      if grain_count > 0:
         grain_size = np.mean(grain_sizes_count)

      else:
         grain_size = 0  # In case no grains were detected

      grain_sizes= []
      grain_sizes.append(grain_size)
        
            

      # Return the results as a dictionary
      return {
        'time': time,
        'grain_sizes': grain_sizes,
        'grain_count': grain_count,
        'center': center,
        'radius': radius
     }


    def preprocess_for_grain_analysis(self, image):
        """Preprocess image for grain detection"""
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        gray = cv2.equalizeHist(gray)
        gaussian_blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        median_blurred = cv2.medianBlur(gaussian_blurred, 5)  # Further reduce noise with median blur
        filtered = cv2.bilateralFilter(median_blurred, d=18, sigmaColor=80, sigmaSpace=80)
        return filtered

    def create_circular_mask(self, image, scale_factor=0.75):
        """Create circular mask for region of interest"""
        height, width = image.shape[:2]
        center = (width // 2, height // 2)
        radius = int(min(center) * scale_factor)
        
        mask = np.zeros((height, width), dtype=np.uint8)
        cv2.circle(mask, center, radius, 255, thickness=-1)
        
        return mask, center, radius

    def detect_grains(self, image, mask):
        """Detect grains using edge detection and contour finding"""
        # Apply Canny edge detection
        edges = cv2.Canny(image, 100,200)
        
        # Dilate edges
        dilated = cv2.dilate(edges, None, iterations=2)
        
        # Apply mask
        masked = cv2.bitwise_and(dilated, dilated, mask=mask)
        
        # Find contours
        contours, _ = cv2.findContours(
            masked, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )
        
        return contours
  
class ThemeManager:
    """Enhanced theme manager with modern color schemes"""
    _instance = None
    _theme = 'light'
    
    THEMES = {
        'light': {
            'bg_primary': '#ffffff',
            'bg_secondary': '#f8f9fa',
            'text_primary': '#212529',
            'text_secondary': '#495057',
            'accent': '#0d6efd',
            'border': '#dee2e6'
        },
        'dark': {
            'bg_primary': '#212529',
            'bg_secondary': '#343a40',
            'text_primary': '#f8f9fa',
            'text_secondary': '#e9ecef',
            'accent': '#0d6efd',
            'border': '#495057'
        }
    }
    
    @classmethod
    def instance(cls):
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance
    
    @classmethod
    def set_theme(cls, theme):
        cls._theme = theme
    
    @classmethod
    def get_theme(cls):
        return cls._theme
    
    @classmethod
    def get_colors(cls, theme=None):
        theme = theme or cls._theme
        return cls.THEMES.get(theme, cls.THEMES['light'])

class ThemeMixin:
    """Enhanced mixin class for theme functionality with improved color schemes"""
    theme_changed = Signal(str)

    def setupTheme(self):
        self.ui.actionLightMode.triggered.connect(lambda: self.setTheme('light'))
        self.ui.actionDarkMode.triggered.connect(lambda: self.setTheme('dark'))
        self.ui.actionSystemMode.triggered.connect(lambda: self.setTheme('system'))
        self.setTheme(ThemeManager.get_theme())

    def setTheme(self, theme):
        ThemeManager.set_theme(theme)
        if hasattr(self.ui, 'actionLightMode'):
            self.ui.actionLightMode.setChecked(theme == 'light')
        if hasattr(self.ui, 'actionDarkMode'):
            self.ui.actionDarkMode.setChecked(theme == 'dark')
        if hasattr(self.ui, 'actionSystemMode'):
            self.ui.actionSystemMode.setChecked(theme == 'system')

        self.applyTheme(theme == 'dark')

    def applyTheme(self, is_dark):
        """Apply enhanced theme styling with improved colors"""
        colors = {
            'dark': {
                'bg_primary': '#1E1E2E',
                'bg_secondary': '#2B2B40',
                'bg_tertiary': '#363650',
                'text_primary': '#FFFFFF',
                'text_secondary': '#B8B8D1',
                'accent': '#7C3AED',
                'accent_hover': '#9061FF',
                'accent_pressed': '#6D28D9',
                'success': '#10B981',
                'warning': '#F59E0B',
                'error': '#EF4444',
                'info': '#3B82F6',
                'border': '#363650',
                'highlight': '#8B5CF6',
                'selection': 'rgba(124, 58, 237, 0.3)',
                'chart_1': '#8B5CF6',
                'chart_2': '#10B981',
                'chart_3': '#3B82F6',
                'chart_4': '#F59E0B',
                'chart_5': '#EC4899'
            },
            'light': {
                'bg_primary': '#FFFFFF',
                'bg_secondary': '#F8FAFC',
                'bg_tertiary': '#F1F5F9',
                'text_primary': '#1E293B',
                'text_secondary': '#64748B',
                'accent': '#7C3AED',
                'accent_hover': '#9061FF',
                'accent_pressed': '#6D28D9',
                'success': '#059669',
                'warning': '#D97706',
                'error': '#DC2626',
                'info': '#2563EB',
                'border': '#E2E8F0',
                'highlight': '#8B5CF6',
                'selection': 'rgba(124, 58, 237, 0.2)',
                'chart_1': '#7C3AED',
                'chart_2': '#059669',
                'chart_3': '#2563EB',
                'chart_4': '#D97706',
                'chart_5': '#DB2777'
            }
        }

        theme = colors['dark'] if is_dark else colors['light']

        style = f"""
            /* Main Window */
            QMainWindow, QDialog {{
                background-color: {theme['bg_primary']};
                color: {theme['text_primary']};
            }}

            /* Menu Bar */
            QMenuBar {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_primary']};
                border-bottom: 1px solid {theme['border']};
                padding: 4px;
            }}

            QMenuBar::item:selected {{
                background-color: {theme['selection']};
                border-radius: 4px;
            }}

            /* Menu */
            QMenu {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_primary']};
                border: 1px solid {theme['border']};
                border-radius: 6px;
                padding: 4px;
            }}

            QMenu::item:selected {{
                background-color: {theme['selection']};
                border-radius: 4px;
            }}

            /* Push Buttons */
            QPushButton {{
                background-color: {theme['accent']};
                color: white;
                border: none;
                border-radius: 6px;
                padding: 8px 16px;
                font-weight: bold;
                min-width: 80px;
                font-size: 14px;
            }}

            QPushButton:hover {{
                background-color: {theme['accent_hover']};
            }}

            QPushButton:pressed {{
                background-color: {theme['accent_pressed']};
            }}

            QPushButton:disabled {{
                background-color: {theme['bg_tertiary']};
                color: {theme['text_secondary']};
            }}

            /* Success Button */
            QPushButton[success="true"] {{
                background-color: {theme['success']};
            }}

            QPushButton[success="true"]:hover {{
                background-color: {theme['success']};
                opacity: 0.9;
            }}

            /* Warning Button */
            QPushButton[warning="true"] {{
                background-color: {theme['warning']};
            }}

            /* Text Inputs */
            QLineEdit, QTextEdit, QPlainTextEdit {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_primary']};
                border: 2px solid {theme['border']};
                border-radius: 6px;
                padding: 8px;
                selection-background-color: {theme['selection']};
                font-size: 14px;
            }}

            QLineEdit:focus, QTextEdit:focus, QPlainTextEdit:focus {{
                border-color: {theme['accent']};
                background-color: {theme['bg_primary']};
            }}

            /* Spin Box */
            QSpinBox, QDoubleSpinBox {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_primary']};
                border: 2px solid {theme['border']};
                border-radius: 6px;
                padding: 6px;
                min-width: 80px;
            }}

            QSpinBox:focus, QDoubleSpinBox:focus {{
                border-color: {theme['accent']};
                background-color: {theme['bg_primary']};
            }}

            /* Progress Bar */
            QProgressBar {{
                background-color: {theme['bg_secondary']};
                border: 2px solid {theme['border']};
                border-radius: 6px;
                height: 20px;
                text-align: center;
                margin: 0px 10px;
            }}

            QProgressBar::chunk {{
                background-color: {theme['accent']};
                border-radius: 4px;
            }}

            /* Labels */
            QLabel {{
                color: {theme['text_primary']};
                font-size: 14px;
            }}

            /* Group Box */
            QGroupBox {{
                background-color: {theme['bg_secondary']};
                border: 2px solid {theme['border']};
                border-radius: 8px;
                margin-top: 12px;
                padding: 15px;
                font-size: 14px;
            }}

            QGroupBox::title {{
                color: {theme['text_primary']};
                subcontrol-origin: margin;
                left: 8px;
                padding: 0 5px;
            }}

            /* Scroll Bars */
            QScrollBar:vertical {{
                background-color: {theme['bg_secondary']};
                width: 14px;
                border-radius: 7px;
                margin: 0px;
            }}

            QScrollBar::handle:vertical {{
                background-color: {theme['accent']};
                min-height: 30px;
                border-radius: 7px;
            }}

            QScrollBar::handle:vertical:hover {{
                background-color: {theme['accent_hover']};
            }}

            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 0px;
            }}

            QScrollBar:horizontal {{
                background-color: {theme['bg_secondary']};
                height: 14px;
                border-radius: 7px;
            }}

            QScrollBar::handle:horizontal {{
                background-color: {theme['accent']};
                min-width: 30px;
                border-radius: 7px;
            }}

            QScrollBar::handle:horizontal:hover {{
                background-color: {theme['accent_hover']};
            }}

            /* Tables */
            QTableView {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_primary']};
                border: 2px solid {theme['border']};
                border-radius: 8px;
                gridline-color: {theme['border']};
                font-size: 14px;
            }}

            QHeaderView::section {{
                background-color: {theme['bg_tertiary']};
                color: {theme['text_primary']};
                padding: 8px;
                border: none;
                border-right: 1px solid {theme['border']};
                border-bottom: 1px solid {theme['border']};
                font-weight: bold;
            }}

            QTableView::item:selected {{
                background-color: {theme['selection']};
                color: {theme['text_primary']};
            }}

            /* Tabs */
            QTabWidget::pane {{
                border: 2px solid {theme['border']};
                border-radius: 8px;
                top: -1px;
            }}

            QTabBar::tab {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_primary']};
                border: 1px solid {theme['border']};
                border-bottom: none;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
                padding: 8px 16px;
                margin-right: 2px;
                font-size: 14px;
            }}

            QTabBar::tab:selected {{
                background-color: {theme['accent']};
                color: white;
                border-bottom: none;
            }}

            /* Status Bar */
            QStatusBar {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_secondary']};
                border-top: 1px solid {theme['border']};
            }}

            /* Tool Tips */
            QToolTip {{
                background-color: {theme['bg_secondary']};
                color: {theme['text_primary']};
                border: 1px solid {theme['border']};
                border-radius: 4px;
                padding: 4px;
            }}
        """
        self.setStyleSheet(style)

class MainWindow(QMainWindow, ThemeMixin):
    def __init__(self):
        super().__init__()
        self.ui = Ui_GRAIN_ANALYSIS_SOFTWARE()
        self.ui.setupUi(self)
        
        # Set window properties
        self.setWindowTitle("Grain Analysis Software")
        self.setMinimumSize(800, 600)
        
        # Initialize theme and UI
        self.setupTheme()
        self.setupUI()
        
        # Initialize components
        self.image_processor = ImageProcessor()
        self.timer = QTimer()
        self.timer.timeout.connect(self.processNextImage)
        
        # Initialize data storage
        self.selected_files = []
        self.processed_results = []
        self.current_index = 0
        self.second_window = None
        
        # Initialize measurement settings
        self.measurement_settings = {
            'magnification': 100.0,
            'unit': 'mm²',
            'scale_bar': 100.0
        }
        

        self.setupCancelButton()
        
    def setupCancelButton(self):
        """Setup cancel button with restart functionality"""
        if hasattr(self.ui, 'cancel_button'):
            self.ui.cancel_button.clicked.connect(self.restartApplication)
            
    def restartApplication(self):
        """Handle application restart"""
        try:
            # Optional: Ask for confirmation
            reply = QMessageBox.question(
                self, 
                'Restart Application',
                'Are you sure you want to restart the application?',
                QMessageBox.Yes | QMessageBox.No, 
                QMessageBox.No
            )
            
            if reply == QMessageBox.Yes:
                # Get the path to the current Python interpreter and script
                program = sys.executable
                script = os.path.abspath(sys.argv[0])
                args = sys.argv[1:]  # Any command line arguments
                
                # Start new instance
                QProcess.startDetached(program, [script] + args)
                
                # Close the current instance
                QApplication.quit()
        except Exception as e:
            QMessageBox.critical(
                self, 
                'Error',
                f'Failed to restart application: {str(e)}'
            )
            
    def setupUI(self):
        """Enhanced UI setup with better spacing and organization"""
        # Setup magnification controls with better layout
        self.setupMagnificationControls()
        
        # Setup buttons with enhanced styling
        self.setupButtons()
        
        # Setup menu actions
        self.setupMenuActions()
        
        # Setup progress tracking with better visual feedback
        self.ui.progressBar.setValue(0)
        self.ui.progressBar.setTextVisible(True)
        self.ui.progressBar.setFormat("%p% Complete")
        self.ui.next_button.setVisible(False)
        
        # Setup status displays with better formatting
        self.ui.textEdit.clear()
        self.ui.textEdit.setReadOnly(True)
        
    def setupButtons(self):
        """Setup main control buttons"""
        self.ui.image_import.clicked.connect(self.getImages)
        self.ui.micronsButton.clicked.connect(self.enableStartProcessing)
        self.ui.process_button.clicked.connect(self.startProcessing)
        self.ui.process_button.setEnabled(False)  # Initially disabled
        self.ui.next_button.clicked.connect(self.openSecondWindow)
    def enableStartProcessing(self):
        """Enable Start Processing button after Save button is clicked"""
        if self.selected_files:
            self.ui.process_button.setEnabled(True)
            self.logStatus("Measurement settings saved. Ready to start processing.")
        
    def setupMenuActions(self):
        """Setup menu bar actions"""
        self.ui.actionSave.triggered.connect(self.saveResults)
        self.ui.actionExit.triggered.connect(self.close)
        
    def setupMagnificationControls(self):
        """Setup magnification controls with better layout"""
        # Setup magnification spinbox with better styling
        self.ui.magnificationSpinBox.setValue(100)
        self.ui.magnificationSpinBox.valueChanged.connect(self.updateMagnification)
        self.ui.magnificationSpinBox.setMinimum(1)
        self.ui.magnificationSpinBox.setMaximum(1000)
        
        # Setup scale bar spinbox
        self.ui.scaleBarSpinBox.setValue(100)
        self.ui.scaleBarSpinBox.valueChanged.connect(self.updateScaleBar)
        self.ui.scaleBarSpinBox.setMinimum(1)
        self.ui.scaleBarSpinBox.setMaximum(1000)
        
        # Setup unit buttons
        self.ui.micronsButton.clicked.connect(lambda: self.setUnit('µm'))
        
    def updateMagnification(self, value):
        """Update magnification setting"""
        self.measurement_settings['magnification'] = float(value)
        
    def updateScaleBar(self, value):
        """Update scale bar measurement setting"""
        self.measurement_settings['scale_bar'] = float(value)
        
    def setUnit(self, unit):
        """Set measurement unit"""
        self.measurement_settings['unit'] = unit
        
        # Update button states
        self.ui.micronsButton.setChecked(unit == 'µm')
        
        # Update status with current scale bar value to show new unit
        current_scale = self.measurement_settings.get('scale_bar', 100)
        current_mag = self.measurement_settings.get('magnification', 100)
        self.logStatus(f"Measurement_overall set to {(current_scale)*(current_mag)}{unit}")
        self.logStatus(f"Scale bar measurement updated to {current_scale} {unit}")
        self.logStatus(f"Magnification updated to {current_mag}{'x'}")
        
    def getImages(self):
        """Open file dialog to select images"""
        file_filter = 'Image Files (*.png *.jpg *.jpeg *.bmp *.gif);;All Files (*)'
        response = QFileDialog.getOpenFileNames(
            parent=self,
            caption='Select Images',
            dir='',
            filter=file_filter
        )
        
        self.selected_files = response[0]
        if self.selected_files:
            self.ui.process_button.setVisible(True)  # Show the Start Processing button
            self.logStatus(f"Selected {len(self.selected_files)} images. Click Save to proceed.")

   
            
    def startProcessing(self):
        """Initialize image processing"""
        if not self.ui.process_button.isEnabled():
            return  # Ensure the button is enabled before proceeding

        self.ui.textEdit.clear()
        
        current_scale = self.measurement_settings.get('scale_bar', 100)
        current_mag = self.measurement_settings.get('magnification', 100)
        self.logStatus(f"Measurement_overall set to {(current_scale)*(current_mag)}{'µm'}")
        self.logStatus(f"Scale bar measurement updated to {current_scale} {'µm'}")
        self.logStatus(f"Magnification updated to {current_mag}{'x'}")

        self.ui.progressBar.setValue(0)
        self.ui.progressBar.setMaximum(len(self.selected_files))

        self.current_index = 0
        self.processed_results = []
        self.timer.start(100)
        
    def processNextImage(self):
        """Process next image in queue"""
        if self.current_index >= len(self.selected_files):
            self.finishProcessing()
            return
            
        try:
            image_path = self.selected_files[self.current_index]
            result = self.image_processor.process_image(image_path)
            
            result['magnification'] = self.measurement_settings['magnification']
            result['scale_bar'] = self.measurement_settings['scale_bar']  # Length of scale bar in micrometers
            true_size =  result['magnification']/((result['scale_bar'] ) **2)# True size of scale bar in micrometers
            pixels_per_um = true_size * 1024
            

            result['grain_sizes'] = [size * pixels_per_um for size in result['grain_sizes']]

            self.processed_results.append(result)

            self.logStatus(
                  f"Processed {os.path.basename(image_path)}:\n"
               f"  Time: {result['time']:.2f}s\n"
              f"  Grains: {result['grain_count']}\n"
            f"  Average Size: {np.mean(result['grain_sizes']):.6f} µm²\n"
        )

        except Exception as e:
            self.logStatus(f"Error processing {image_path}: {str(e)}")
            
        self.current_index += 1
        self.ui.progressBar.setValue(self.current_index)
        
    def finishProcessing(self):
        """Cleanup after processing complete"""
        self.timer.stop()
        self.logStatus("Processing complete!")
        self.ui.next_button.setVisible(True)
        
    def openSecondWindow(self):
        """Open analysis window"""
        if not self.second_window:
            self.second_window = SecondWindow(self)
            self.theme_changed.connect(self.second_window.setTheme)
        
        # Pass measurement settings to second window
        self.second_window.updateResults(self.processed_results)
        self.hide()
        self.second_window.show()
        
    def saveResults(self):
        """Save processing results to CSV"""
        if not self.processed_results:
            self.logStatus("No results to save!")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Results", "", "CSV Files (*.csv)"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow([
                        'Time (s)',
                        'Grain Count',
                        f'Average Grain Size(µm²) ({self.measurement_settings})',
                     
                    ])
                    
                    for result in self.processed_results:
                     
                        if (result['grain_count'] > 3 and 
                            result['grain_sizes'] and 
                            np.mean(result['grain_sizes']) > 0 and 
                            0 < result['time'] <= 6000):
                            writer.writerow([
                                f"{result['time']:.2f}",
                                result['grain_count'],
                                f"{np.mean(result['grain_sizes']):.2f}"
                            ])
                
                self.logStatus(f"Results saved to {file_path}")
                
            except Exception as e:
                self.logStatus(f"Error saving results: {str(e)}")
                
    def logStatus(self, message):
        """Add message to status log"""
        self.ui.textEdit.append(message)
        
    def closeEvent(self, event):
        """Handle window close event"""
        if self.second_window:
            self.second_window.close()
        event.accept()

class SecondWindow(QMainWindow, ThemeMixin):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_secondwindow()
        self.ui.setupUi(self)
        self.setWindowIcon(parent.windowIcon())

        # Initialize
        self.setupTheme()
        self.setupUI()
        self.parent = parent
        self.results = []
        self.setupMenuActions()
        
        self.ui.actionExit.triggered.disconnect()  # Disconnect default close action
        self.ui.actionExit.triggered.connect(self.goBack)  # Connect to new back function
        
        # Create layouts for plots
        self.setupPlotLayouts()
    def closeEvent(self, event):
        """Override close event to implement back behavior"""
        self.goBack()
        event.ignore()  # Prevent actual window closure
        
    def goBack(self):
        """Return to main window"""
        if self.parent:
            self.parent.show()  # Show main window
            self.hide()  # Hide second window instead of closing it    
    def setupUI(self):
        """Setup UI connections and initial state"""
        # Setup buttons
        self.ui.analyzepush.clicked.connect(self.generateReport)
        self.ui.import_to_csv.clicked.connect(self.exportToCsv)
        self.ui.print_image_data.clicked.connect(self.printData)
        
        # Setup text display
        self.ui.textEdit.setReadOnly(True)
        self.ui.textEdit.setLineWrapMode(QTextEdit.WidgetWidth)
     
    def setupMenuActions(self):
        """Setup menu bar actions"""
        self.ui.actionSave.triggered.connect(self.exportToCsv)
        self.ui.actionExit.triggered.connect(self.close)
        
    def setupPlotLayouts(self):
        """Setup layouts for interactive plots"""
        # Layout for grain growth plot
        self.plot_layout1 = QVBoxLayout()
        self.plot_widget1 = QWidget()
        self.plot_widget1.setLayout(self.plot_layout1)
        
        # Create a QScrollArea for the first plot
        scroll1 = QScrollArea()
        scroll1.setWidget(self.plot_widget1)
        scroll1.setWidgetResizable(True)
        self.ui.Grain_growth_2.setWidget(scroll1)
        
        # Layout for grain count plot
        self.plot_layout2 = QVBoxLayout()
        self.plot_widget2 = QWidget()
        self.plot_widget2.setLayout(self.plot_layout2)
        
        # Create a QScrollArea for the second plot
        scroll2 = QScrollArea()
        scroll2.setWidget(self.plot_widget2)
        scroll2.setWidgetResizable(True)
        self.ui.Change_in_GB_no.setWidget(scroll2)
    
    def updateResults(self, results):
        """Update with new processing results"""
        try:
            # Get magnification from parent window
            magnification = self.parent.ui.magnificationSpinBox.value()
            
            # Load and process grain data
            self.results = self.load_grain_data(results, magnification)
            self.updateVisualizations()
            self.display_velocity()
        except Exception as e:
            self.ui.textEdit.setText(f"Error updating results: {str(e)}")

    def load_grain_data(self, images, magnification):
        """Load and validate grain data from images"""
        results = []
        
        for img_result in images:
            # Get grain data from image analysis results
            result = {
                'time': img_result['time'],
                'grain_count': img_result['grain_count'],
                'grain_sizes': img_result['grain_sizes'],
                'magnification': magnification,
                'unit': img_result.get('unit', 'μm²')
            }
            results.append(result)
            
        return results
    
    def generateReport(self):
     doc = Document()
    
     # Add title
     title = doc.add_heading('Grain Analysis Report', 0)
     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
     
     
    
     # Add Grain Growth Plot
     doc.add_heading('Grain Growth Analysis', level=1)
     filtered_results = [r for r in self.results if 0 < r['time'] < 6000]
     times = [r['time'] for r in filtered_results]
     sizes = [np.mean(r['grain_sizes']) if r['grain_sizes'] else 0 for r in filtered_results]
    
     plt.figure(figsize=(8, 6))
     plt.scatter(times, sizes, color='blue')
    
     if len(times) > 1:
        from sklearn.linear_model import LinearRegression
        model = LinearRegression()
        X = np.array(times).reshape(-1, 1)
        model.fit(X, sizes)
        line_x = np.linspace(min(times), max(times), 100)
        line_y = model.predict(line_x.reshape(-1, 1))
        plt.plot(line_x, line_y, 'r-', label='Trend')
    
     plt.xlabel('Time (s)')
     plt.ylabel('Average Grain Size')
     plt.title('Grain Growth vs Time')
     plt.grid(True)
    
     # Save plot to memory
     img_stream = BytesIO()
     plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
     img_stream.seek(0)
     doc.add_picture(img_stream, width=Inches(6))
     plt.close()
    
     doc.add_paragraph('')  # Add spacing
    
    # Add Grain Count Plot
     doc.add_heading('Grain Count Analysis', level=1)
     counts = [r['grain_count'] for r in filtered_results]
    
     plt.figure(figsize=(8, 6))
     plt.scatter(times, counts, color='green')
    
     if len(times) > 1:
        model = LinearRegression()
        X = np.array(times).reshape(-1, 1)
        model.fit(X, counts)
        line_x = np.linspace(min(times), max(times), 100)
        line_y = model.predict(line_x.reshape(-1, 1))
        plt.plot(line_x, line_y, 'r-', label='Trend')
    
     plt.xlabel('Time (s)')
     plt.ylabel('Grain Count')
     plt.title('Grain Count vs Time')
     plt.grid(True)
    
     # Save plot to memory
     img_stream = BytesIO()
     plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
     img_stream.seek(0)
     doc.add_picture(img_stream, width=Inches(6))
     plt.close()
    
    # Save document
     file_path, _ = QFileDialog.getSaveFileName(
        self,
        "Save Report",
        "",
        "Word Documents (*.docx)"
    )
    
     if file_path:
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        doc.save(file_path)

    
    def updateVisualizations(self):
        """Update all visualization plots"""
        if not self.results:
            self.ui.textEdit.setText("No results available for visualization.")
            return
            
        self.plotGrainGrowth()
        self.plotGrainCount()

    def calculate_grain_boundary_velocity(self):
        """Calculate grain boundary velocity using strategic sampling of valid images"""
        # First filter for valid images only
        valid_results = [
            result for result in self.results
            if (result['time'] > 0 and 
                result['grain_count'] > 3 and 
                result['grain_sizes'] and 
                np.mean(result['grain_sizes']) > 0 and 
                result['time'] <= 6000)
        ]
        
        if len(valid_results) < 2:
            raise ValueError("Need at least 2 valid images to calculate velocity")
            
        # Get indices for sampling (first, quarter, half, last)
        n = len(valid_results)
        sample_points = [
            valid_results[0],          # First valid image
            valid_results[n // 4],     # Quarter way
            valid_results[n // 2],     # Halfway
            valid_results[n - 1]       # Last valid image
        ]
        
        velocities = []
        for i in range(1, len(sample_points)):
            prev_result = sample_points[i-1]
            curr_result = sample_points[i]
            
            # Calculate grains per unit area for both points
            def calculate_grains_per_area(result):
                area_mm2 = np.mean(result['grain_sizes'])/1000000  # Convert to mm²
                magnification = result['magnification']
                grains_per_area = (magnification*magnification) / (area_mm2) * result['grain_count']
                
                # Calculate diameter using empirical formula: D = 998.3 * A^(-0.499)
                diameter = 998.3 * (grains_per_area ** -0.499) * 1000  # Convert to microns
                return diameter
                
            prev_diameter = calculate_grains_per_area(prev_result)
            curr_diameter = calculate_grains_per_area(curr_result)
            
            # Calculate velocity
            time_diff = curr_result['time'] - prev_result['time']
            diameter_change = abs(curr_diameter - prev_diameter)
            velocity = diameter_change / time_diff
            velocities.append(velocity)
                
        if not velocities:
            raise ValueError("Could not calculate valid velocities from the data")
            
        return np.mean(velocities)

    def display_velocity(self):
        """Display grain boundary velocity in the UI with μm/s units"""
        try:
            velocity = self.calculate_grain_boundary_velocity()
            
            velocity_text = (
                f"Average velocity: {velocity:.2f} μm/s\n"
                f"Analysis based on {len(self.results)} images\n"
                f"Magnification: {self.parent.ui.magnificationSpinBox.value()}\n"
            )
            
            self.ui.textEdit.append(velocity_text)
            
        except ValueError as e:
            error_text = f"Error calculating velocity: {str(e)}"
            self.ui.textEdit.append(error_text)

    def plotGrainGrowth(self):
        """Plot grain growth over time with enhanced interactivity"""
        # Filter out zero values and limit time range
        valid_data = [(r['time'], np.mean(r['grain_sizes'])) 
                     for r in self.results 
                     if r['grain_sizes'] and 
                     np.mean(r['grain_sizes']) > 0 and 
                     r['time'] <= 6000 and r['time'] > 1]
        
        if not valid_data:
            self.ui.textEdit.append("No valid grain size data available.")
            return
            
        times, sizes = zip(*valid_data)
        
        fig = Figure(figsize=(8, 6), dpi=100)
        ax = fig.add_subplot(111)
        scatter = ax.scatter(times, sizes, color='blue', picker=True, 
                           label='Measured Data')
        
        # Add trend line
        if len(times) > 1:
            model = LinearRegression()
            X = np.array(times).reshape(-1, 1)
            y = np.array(sizes)
            model.fit(X, y)
            line_x = np.linspace(min(times), max(times), 100)
            line_y = model.predict(line_x.reshape(-1, 1))
            r2 = model.score(X, y)
            ax.plot(line_x, line_y, 'r-', label=f'Trend Line (R² = {r2:.3f})')
        
        ax.set_xlabel('Time (s)', fontsize=12)
        ax.set_ylabel(f'Average Grain Size ({self.results[0]["unit"]})', 
                     fontsize=12)
        ax.set_title('Grain Growth vs Time', fontsize=14, pad=20)
        ax.grid(True, linestyle='--', alpha=0.7)
        ax.legend(fontsize=10)
        fig.tight_layout()

        # Clear previous plot
        for i in reversed(range(self.plot_layout1.count())): 
            self.plot_layout1.itemAt(i).widget().setParent(None)
        
        # Add new plot with navigation toolbar
        canvas = FigureCanvas(fig)
        toolbar = NavigationToolbar2QT(canvas, self)
        for action in toolbar.actions():
            if action.text() == 'Close':
                toolbar.removeAction(action)
                break
        self.plot_layout1.addWidget(toolbar)
        self.plot_layout1.addWidget(canvas)
        
        def on_pick(event):
            ind = event.ind[0]
            self.ui.textEdit.append(
                f"\nGrain Growth Data Point:\n"
                f"Time: {times[ind]:.2f} s\n"
                f"Average Size: {sizes[ind]:.2f} {self.results[0]['unit']}\n"
            )
        
        canvas.mpl_connect('pick_event', on_pick)

    def plotGrainCount(self):
        """Plot grain count over time with enhanced interactivity"""
        # Filter out zero values and limit time range
        valid_data = [(r['time'], r['grain_count']) 
                     for r in self.results 
                     if r['grain_count'] > 3 and 
                     r['time'] <= 6000 and r['time'] > 1]
        
        if not valid_data:
            self.ui.textEdit.append("No valid grain count data available.")
            return
            
        times, counts = zip(*valid_data)
        
        fig = Figure(figsize=(8, 6), dpi=100)
        ax = fig.add_subplot(111)
        scatter = ax.scatter(times, counts, color='green', picker=True,
                           label='Measured Data')
        
        # Add trend line
        if len(times) > 1:
            model = LinearRegression()
            X = np.array(times).reshape(-1, 1)
            y = np.array(counts)
            model.fit(X, y)
            line_x = np.linspace(min(times), max(times), 100)
            line_y = model.predict(line_x.reshape(-1, 1))
            r2 = model.score(X, y)
            ax.plot(line_x, line_y, 'r-', label=f'Trend Line (R² = {r2:.3f})')
        
        ax.set_xlabel('Time (s)', fontsize=12)
        ax.set_ylabel('Grain Count', fontsize=12)
        ax.set_title('Grain Count vs Time', fontsize=14, pad=20)
        ax.grid(True, linestyle='--', alpha=0.7)
        ax.legend(fontsize=10)
        fig.tight_layout()

        # Clear previous plot
        for i in reversed(range(self.plot_layout2.count())): 
            self.plot_layout2.itemAt(i).widget().setParent(None)
        
        # Add new plot with navigation toolbar
        canvas = FigureCanvas(fig)
        toolbar = NavigationToolbar2QT(canvas, self)
        for action in toolbar.actions():
            if action.text() == 'Close':
                toolbar.removeAction(action)
                break
        self.plot_layout2.addWidget(toolbar)
        self.plot_layout2.addWidget(canvas)
        
        def on_pick(event):
            ind = event.ind[0]
            self.ui.textEdit.append(
                f"\nGrain Count Data Point:\n"
                f"Time: {times[ind]:.2f} s\n"
                f"Grain Count: {counts[ind]}\n"
            )
        
        canvas.mpl_connect('pick_event', on_pick)

    def exportToCsv(self):
        """Export results to CSV with non-zero values only"""
        if not self.results:
            self.ui.textEdit.append("No results available to export.")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Results", "", "CSV Files (*.csv)"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow([
                        'Time (s)', 
                        'Grain Count',
                        f'Average Grain Size(µm²) ({self.results[0]})'
                    ])
                    
                    for result in self.results:
                        # Only write non-zero values
                        if (result['grain_count'] > 3 and 
                            result['grain_sizes'] and 
                            np.mean(result['grain_sizes']) > 0 and 
                            0 < result['time'] <= 6000):
                            writer.writerow([
                                f"{result['time']:.2f}",
                                result['grain_count'],
                                f"{np.mean(result['grain_sizes']):.2f}"
                            ])
                
                self.ui.textEdit.append(f"Results successfully exported to {file_path}")
                
            except Exception as e:
                self.ui.textEdit.append(f"Error exporting results: {str(e)}")

    def printData(self):
        """Print current results to text display, excluding zero values"""
        if not self.results:
            self.ui.textEdit.setText("No results available.")
            return
            
        text = "Analysis Results (Non-zero values only):\n\n"
        for i, result in enumerate(self.results, 1):
            # Skip if both grain count and average size are zero
            if (result['grain_count'] <= 0 or 
                not result['grain_sizes'] or 
                np.mean(result['grain_sizes']) <= 0):
                continue
                
            text += f"Image {i}:\n"
            text += f"Time: {result['time']:.2f} s\n"
            text += f"Grain Count: {result['grain_count']}\n"
            text += (f"Average Grain Size: {np.mean(result['grain_sizes']):.2f} "
                    f"{result['unit']}\n\n")
        
        self.ui.textEdit.setText(text)

    def setupPlotLayouts(self):
        """Setup layouts for interactive plots"""
        # Replace QGraphicsView with QWidget for the first plot
        self.plot_widget1 = QWidget()
        self.plot_layout1 = QVBoxLayout(self.plot_widget1)
        
        # Create a scroll area for the first plot
        scroll1 = QScrollArea()
        scroll1.setWidget(self.plot_widget1)
        scroll1.setWidgetResizable(True)
        
        # Replace the QGraphicsView with our scroll area
        if hasattr(self.ui, 'Grain_growth_2'):
            # Remove the existing QGraphicsView
            old_widget = self.ui.Grain_growth_2
            parent_layout = old_widget.parent().layout()
            if parent_layout:
                parent_layout.replaceWidget(old_widget, scroll1)
            old_widget.hide()
            old_widget.deleteLater()
        self.ui.Grain_growth_2 = scroll1
        
        # Repeat for the second plot
        self.plot_widget2 = QWidget()
        self.plot_layout2 = QVBoxLayout(self.plot_widget2)
        
        # Create a scroll area for the second plot
        scroll2 = QScrollArea()
        scroll2.setWidget(self.plot_widget2)
        scroll2.setWidgetResizable(True)
        
        # Replace the QGraphicsView with our scroll area
        if hasattr(self.ui, 'Change_in_GB_no'):
            # Remove the existing QGraphicsView
            old_widget = self.ui.Change_in_GB_no
            parent_layout = old_widget.parent().layout()
            if parent_layout:
                parent_layout.replaceWidget(old_widget, scroll2)
            old_widget.hide()
            old_widget.deleteLater()
        self.ui.Change_in_GB_no = scroll2

if __name__ == '__main__':
    app = QApplication(sys.argv)
    
    # Set application style
    app.setStyle('Fusion')
    
    # Initialize theme manager
    theme_manager = ThemeManager.instance()
    
    # Create and show main window
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec())